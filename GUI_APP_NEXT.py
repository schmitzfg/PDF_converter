import tkinter as tk
from tkinter import Label
from docx import Document
from PIL import Image, ImageTk
from tkinter.filedialog import askopenfile, asksaveasfilename
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
import fitz
from docx.shared import Inches
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont


pdfmetrics.registerFont(TTFont("DejaVu", "DejaVuSans.ttf"))
## pdfminer.six  library PDF

BG = "#f7ffde"

root = tk.Tk()
##--------app size and position---------------
width = 950 
height = 600
screen_width = root.winfo_screenwidth()
screen_height = root.winfo_screenheight()
x = (screen_width - width) // 2
y = (screen_height - height) // 2
root.geometry(f"{width}x{height}+{x}+{y}")
##-------------------------------------------------

root.title('PDF text extract')
root.config(padx=20, pady=20, bg=BG, highlightthickness=0)   # highlightthickness=0 - border


# -------- MAIN CONTAINER ----------
frame = tk.Frame(root, bg=BG)
frame.pack(fill="both", expand=True)


# -------- LEFT SIDE (LOGO) ----------
left_frame = tk.Frame(frame, bg=BG)
left_frame.pack(side="left", fill="both", expand=True)


# logo
logo = Image.open('logo.png')
logo = ImageTk.PhotoImage(logo)       # convert in tkinter image

logo_label = tk.Label(left_frame, image=logo, bg=BG)
logo_label.image = logo
logo_label.pack(pady=40)


label = Label(left_frame,
              text="✍️ 📃",
              font=("Raleway",20,"bold"),
              bg=BG,
              fg="#ff9800",
              compound = "bottom")

label.pack(pady=70)



def reset_ui():
    """
    modify Label color and text
    """
    label.config(text="✍️ 📃", fg="#ff9800")



##----------- middle frame ----------
middle_frame = tk.Frame(frame, bg=BG)
middle_frame.pack(side="left", fill="both", expand=True)


# -------- RIGHT SIDE (CONTROLS) ----------
right_frame = tk.Frame(frame, bg=BG)
right_frame.pack(side="right", fill="both", expand=True)

# Instructions
instructions = tk.Label(right_frame, text="Select a PDF file", font=("Raleway", 18), bg=BG)
instructions.pack(pady=(80, 20))


browse_text = tk.StringVar(value="Browse")
current_file_type = None

##===================================================================================



def open_file():
    global current_file_type, pdf_blocks, pdf_images
    browse_text.set("loading...")

    file = askopenfile(
        parent=root,
        mode='rb',
        title='Choose a file',
        filetypes=[("PDF file", "*.pdf"), ("Word file", "*.docx")]
    )

    if not file:
        return

    filename = file.name
    content = ""

    pdf_blocks = []
    pdf_images = []

    # ================= PDF =================
    if filename.lower().endswith(".pdf"):
        current_file_type = "pdf"
        save_btn.config(text="Save as DOCX")

        doc = fitz.open(filename)

        for page_num, page in enumerate(doc):

            # 🔹 TEXT BLOCKS (layout)
            blocks = page.get_text("blocks")
            blocks.sort(key=lambda b: (b[1], b[0]))

            for b in blocks:
                text = b[4].strip()
                if text:
                    pdf_blocks.append((page_num, b, text))
                    content += text + "\n"

            # 🔹 IMAGES
            images = page.get_images(full=True)
            for img_index, img in enumerate(images):
                xref = img[0]
                base_image = doc.extract_image(xref)
                image_bytes = base_image["image"]

                img_name = f"img_{page_num}_{img_index}.png"
                with open(img_name, "wb") as f:
                    f.write(image_bytes)

                pdf_images.append(img_name)

    # ================= DOCX =================
    elif filename.lower().endswith(".docx"):
        current_file_type = "docx"
        save_btn.config(text="Save as PDF")

        doc = Document(filename)

        for para in doc.paragraphs:
            content += para.text + "\n"

    # ================= TEXTBOX =================
    text_box.delete("1.0", tk.END)
    text_box.insert("1.0", content)

    browse_text.set("Browse")



def handle_save():
    """
    Handle the which fucntion should use button save
    """
    if current_file_type == "pdf":
        save_in_DOCX()
    elif current_file_type == "docx":
        save_in_PDF()
    else:
        label.config(text="❌ No file loaded", fg="red")
        print("❌ No file loaded")
    root.after(5000, reset_ui)



## ----------- browse button ---------------
browse_btn = tk.Button(right_frame, 
                       command=open_file,
                       textvariable=browse_text, 
                       font=("Raleway", 14), 
                       bg="#20bebe", 
                       fg=BG, 
                       height=2,
                       width=15,
                       activebackground="#ff9800",
                       activeforeground= "#07ea29",
                       )
browse_text.set("Browse")
browse_btn.pack(pady=10)



# -------- TEXT OUTPUT ----------
text_box = tk.Text(
    middle_frame,
    height=45,
    width=50,
    padx=10,
    pady=10
)
text_box.pack(pady=20)



## -------- add save icon -------------
# save_img = Image.open("save_icon.png")
# save_img = save_img.resize((30, 30))
# save_img = ImageTk.PhotoImage(save_img)


            
def save_in_DOCX():
    if current_file_type != "pdf":
        return

    file_path = asksaveasfilename(
        defaultextension=".docx",
        filetypes=[("Word file", "*.docx")]
    )

    if not file_path:
        return

    docx = Document()

    # ================= COLOANE =================
    left_col = []
    right_col = []

    for page_num, b, text in pdf_blocks:
        x0 = b[0]
        page_width = 600  # aproximare

        if x0 < page_width / 2:
            left_col.append(text)
        else:
            right_col.append(text)

    table = docx.add_table(rows=1, cols=2)
    cell_left = table.rows[0].cells[0]
    cell_right = table.rows[0].cells[1]

    for line in left_col:
        cell_left.add_paragraph(line)

    for line in right_col:
        cell_right.add_paragraph(line)

    # ================= IMAGES =================
    for img in pdf_images:
        try:
            docx.add_picture(img, width=Inches(3))
        except:
            pass

    docx.save(file_path)
            
    label.config(text="✅ Document saved", fg="green")
    root.after(5000, reset_ui)






def save_in_PDF():
    content = text_box.get("1.0", tk.END).strip()

    if not content:
        return

    file_path = asksaveasfilename(
                                defaultextension=".pdf",
                                filetypes=[("PDF file", "*.pdf")])

    if not file_path:
        return

    doc = SimpleDocTemplate(file_path)
    styles = getSampleStyleSheet()

    style = styles["Normal"]
    style.fontName = "DejaVu"
    # style.fontSize = 11
    # style.leading = 14

    story = []

    for line in content.split("\n"):
        story.append(Paragraph(line, styles["Normal"]))

    doc.build(story)
    label.config(text="✅ Document saved", fg="green")
    root.after(5000, reset_ui)



##---------Save btn----------
save_btn = tk.Button(
    right_frame,
    text="SAVE",
    command=handle_save,
    # image=save_img,
    compound="left",
    font=("Raleway", 14),
    bg="#ff9800",
    fg=BG,
    width=15,
    height=2,
    activebackground="#20bebe",
    activeforeground= "#07ea29"
)
# save_btn.image = save_img
save_btn.pack(pady=10)




root.mainloop()